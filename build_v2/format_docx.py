from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.enum.style import WD_STYLE_TYPE

IN='latest_thesis_raw.docx'
OUT='latest_thesis_final.docx'
doc=Document(IN)

# Page setup: A4 + public-document-like margins
for sec in doc.sections:
    sec.top_margin = Cm(3.7)
    sec.bottom_margin = Cm(3.5)
    sec.left_margin = Cm(2.8)
    sec.right_margin = Cm(2.6)
    sec.header_distance = Cm(1.5)
    sec.footer_distance = Cm(1.75)

# Helpers
def set_run_font(run, cn='FangSong', en='Times New Roman', size=16, bold=None):
    run.font.name = en
    run.font.size = Pt(size)
    if bold is not None:
        run.font.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.insert(0, rfonts)
    rfonts.set(qn('w:eastAsia'), cn)
    rfonts.set(qn('w:ascii'), en)
    rfonts.set(qn('w:hAnsi'), en)

def style_paragraph(p, cn='FangSong', en='Times New Roman', size=16, bold=None, first_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=28, before=0, after=0):
    p.alignment = align
    pf = p.paragraph_format
    pf.space_before = Pt(before)
    pf.space_after = Pt(after)
    pf.line_spacing = Pt(line)
    pf.first_line_indent = Pt(size*2) if first_indent else Pt(0)
    for r in p.runs:
        set_run_font(r, cn, en, size, bold)

# Base styles
styles=doc.styles
normal=styles['Normal']
normal.font.name='Times New Roman'; normal.font.size=Pt(16)
normal._element.rPr.rFonts.set(qn('w:eastAsia'),'FangSong')
normal.paragraph_format.line_spacing=Pt(28)
normal.paragraph_format.first_line_indent=Pt(32)

for name in ['Heading 1','Heading 2','Heading 3']:
    if name not in styles:
        styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)

h1=styles['Heading 1']; h1.font.name='SimHei'; h1.font.size=Pt(18); h1.font.bold=False
h1._element.rPr.rFonts.set(qn('w:eastAsia'),'SimHei')
h1.paragraph_format.alignment=WD_ALIGN_PARAGRAPH.CENTER
h1.paragraph_format.space_before=Pt(18); h1.paragraph_format.space_after=Pt(12); h1.paragraph_format.line_spacing=Pt(30); h1.paragraph_format.page_break_before=True

h2=styles['Heading 2']; h2.font.name='KaiTi'; h2.font.size=Pt(16); h2.font.bold=False
h2._element.rPr.rFonts.set(qn('w:eastAsia'),'KaiTi')
h2.paragraph_format.space_before=Pt(12); h2.paragraph_format.space_after=Pt(6); h2.paragraph_format.line_spacing=Pt(28); h2.paragraph_format.keep_with_next=True

h3=styles['Heading 3']; h3.font.name='FangSong'; h3.font.size=Pt(16); h3.font.bold=True
h3._element.rPr.rFonts.set(qn('w:eastAsia'),'FangSong')
h3.paragraph_format.space_before=Pt(8); h3.paragraph_format.space_after=Pt(4); h3.paragraph_format.line_spacing=Pt(28); h3.paragraph_format.keep_with_next=True

# Page break and TOC markers
for p in list(doc.paragraphs):
    txt=p.text.strip()
    if txt=='[PAGEBREAK]':
        p.text=''
        p.add_run().add_break()
        # convert run break to page break
        br=p.runs[-1]._element.find(qn('w:br'))
        if br is not None: br.set(qn('w:type'),'page')
    elif txt=='[TOC_FIELD]':
        p.text=''
        p.alignment=WD_ALIGN_PARAGRAPH.LEFT
        r=p.add_run()
        fldChar=OxmlElement('w:fldChar'); fldChar.set(qn('w:fldCharType'),'begin')
        instrText=OxmlElement('w:instrText'); instrText.set(qn('xml:space'),'preserve'); instrText.text=' TOC \\o "1-3" \\h \\z \\u '
        fldChar2=OxmlElement('w:fldChar'); fldChar2.set(qn('w:fldCharType'),'separate')
        t=OxmlElement('w:t'); t.text='目录将在 Word 中自动更新'
        fldChar2.append(t)
        fldChar3=OxmlElement('w:fldChar'); fldChar3.set(qn('w:fldCharType'),'end')
        r._r.append(fldChar); r._r.append(instrText); r._r.append(fldChar2); r._r.append(fldChar3)

# Update fields on open
settings=doc.settings._element
upd=OxmlElement('w:updateFields'); upd.set(qn('w:val'),'true'); settings.append(upd)

# Paragraph-specific styling
in_abstract=False; in_en_abstract=False; in_refs=False
for idx,p in enumerate(doc.paragraphs):
    t=p.text.strip()
    if not t: continue
    if t=='分散经营与组织化运营':
        style_paragraph(p, cn='STZhongsong', size=26, bold=True, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=40, before=100, after=8)
    elif t=='——普洱旅居型村庄经营秩序的生成逻辑':
        style_paragraph(p, cn='STZhongsong', size=20, bold=False, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=34, before=0, after=60)
    elif t.startswith('硕士学位论文（田野深化整合版）') or t.startswith('专业名称：') or t.startswith('研究方向：') or t=='二〇二六年八月':
        style_paragraph(p, cn='FangSong', size=16, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=30, before=0, after=6)
    elif t.startswith('From Dispersed Operation') or t.startswith('The Formation Logic') or t.startswith('Candidate:') or t.startswith('Major:') or t.startswith('Supervisor:'):
        style_paragraph(p, cn='Times New Roman', en='Times New Roman', size=14, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=26, before=6, after=6)
    elif t=='摘 要':
        in_abstract=True; in_en_abstract=False; in_refs=False
        style_paragraph(p, cn='SimHei', size=18, bold=False, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=30, before=0, after=12)
    elif t=='ABSTRACT':
        in_abstract=False; in_en_abstract=True; in_refs=False
        style_paragraph(p, cn='Times New Roman', en='Times New Roman', size=16, bold=True, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=28, before=0, after=12)
    elif t=='目 录':
        in_abstract=False; in_en_abstract=False; in_refs=False
        style_paragraph(p, cn='SimHei', size=18, bold=False, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=30, before=0, after=12)
    elif t=='参考文献':
        in_abstract=False; in_en_abstract=False; in_refs=True
        style_paragraph(p, cn='SimHei', size=18, bold=False, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=30, before=0, after=12)
    elif t=='致 谢':
        in_refs=False
        style_paragraph(p, cn='SimHei', size=18, bold=False, first_indent=False, align=WD_ALIGN_PARAGRAPH.CENTER, line=30, before=0, after=12)
    elif in_en_abstract:
        style_paragraph(p, cn='Times New Roman', en='Times New Roman', size=12, first_indent=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=22, before=0, after=6)
    elif in_refs:
        style_paragraph(p, cn='SimSun', size=12, first_indent=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=22, before=0, after=3)
    elif p.style.name.startswith('Heading'):
        # preserve heading style fonts
        for r in p.runs:
            if p.style.name=='Heading 1': set_run_font(r,'SimHei','Times New Roman',18,False)
            elif p.style.name=='Heading 2': set_run_font(r,'KaiTi','Times New Roman',16,False)
            else: set_run_font(r,'FangSong','Times New Roman',16,True)
    elif t.startswith('关键词：'):
        style_paragraph(p, cn='FangSong', size=16, first_indent=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=28, before=8, after=0)
    else:
        style_paragraph(p, cn='FangSong', size=16, first_indent=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY, line=28)

# Tables
for table in doc.tables:
    table.alignment=WD_TABLE_ALIGNMENT.CENTER
    table.autofit=True
    for row in table.rows:
        for cell in row.cells:
            cell.vertical_alignment=WD_CELL_VERTICAL_ALIGNMENT.CENTER
            tcPr=cell._tc.get_or_add_tcPr()
            tcBorders=tcPr.first_child_found_in('w:tcBorders')
            if tcBorders is None:
                tcBorders=OxmlElement('w:tcBorders'); tcPr.append(tcBorders)
            for edge in ('top','left','bottom','right','insideH','insideV'):
                tag='w:'+edge
                el=tcBorders.find(qn(tag))
                if el is None:
                    el=OxmlElement(tag); tcBorders.append(el)
                el.set(qn('w:val'),'single'); el.set(qn('w:sz'),'4'); el.set(qn('w:color'),'000000')
            for p in cell.paragraphs:
                p.paragraph_format.first_line_indent=Pt(0)
                p.paragraph_format.line_spacing=Pt(20)
                for r in p.runs: set_run_font(r,'SimSun','Times New Roman',10.5,None)
    # bold first row
    if table.rows:
        for cell in table.rows[0].cells:
            for p in cell.paragraphs:
                for r in p.runs: r.font.bold=True

# Header/footer
for sec in doc.sections:
    hp=sec.header.paragraphs[0]
    hp.text='武汉大学硕士学位论文'
    hp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    for r in hp.runs: set_run_font(r,'SimSun','Times New Roman',9,False)
    fp=sec.footer.paragraphs[0]
    fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=fp.add_run()
    begin=OxmlElement('w:fldChar'); begin.set(qn('w:fldCharType'),'begin')
    instr=OxmlElement('w:instrText'); instr.set(qn('xml:space'),'preserve'); instr.text=' PAGE '
    sep=OxmlElement('w:fldChar'); sep.set(qn('w:fldCharType'),'separate')
    end=OxmlElement('w:fldChar'); end.set(qn('w:fldCharType'),'end')
    r._r.append(begin); r._r.append(instr); r._r.append(sep); r._r.append(end)
    set_run_font(r,'SimSun','Times New Roman',9,False)

# Remove header/footer on first page by first-page different setting
for sec in doc.sections:
    sec.different_first_page_header_footer=True

# Keep headings with following content
for p in doc.paragraphs:
    if p.style.name in ('Heading 1','Heading 2','Heading 3'):
        p.paragraph_format.keep_with_next=True
        p.paragraph_format.widow_control=True

# Document metadata
doc.core_properties.title='分散经营与组织化运营——普洱旅居型村庄经营秩序的生成逻辑'
doc.core_properties.subject='武汉大学社会学硕士学位论文田野深化整合稿'
doc.core_properties.author=''

doc.save(OUT)
print(OUT)
